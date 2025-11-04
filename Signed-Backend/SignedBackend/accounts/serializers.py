from rest_framework import serializers
from .models import User, EmployerProfile, ApplicantProfile, Company
import os
import fitz
import numpy as np
from docx import Document
import mimetypes
import subprocess, os

from sentence_transformers import SentenceTransformer
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")

class UserSerializer(serializers.ModelSerializer):
    class Meta:
        model = User
        fields = ["id", "email", "first_name", "last_name", "role", "firebase_uid"]

class CompanySerializer(serializers.ModelSerializer):
    class Meta:
        model = Company
        fields = ["id", "name", "size", "website", "logo"]

class EmployerProfileSerializer(serializers.ModelSerializer):
    company = CompanySerializer(read_only=True)
    
    class Meta:
        model = EmployerProfile
        fields = ["company", "job_title", "profile_image", "bio", "location"]

    def update(self, instance, validated_data):
        company_data = validated_data.pop("company", None)

        # to update company fields
        if company_data:
            company = instance.company
            for attr, value in company_data.items():
                setattr(company, attr, value)
            company.save()

        # to update employer profile fields
        return super().update(instance, validated_data)
        
class ApplicantProfileSerializer(serializers.ModelSerializer):
    class Meta:
        model = ApplicantProfile
        fields = ["major", "school", "bio", "resume", "resume_file", "skills", "portfolio_url", "profile_image", "vector_embedding", "personality_type"]
        
class MeSerializer(serializers.ModelSerializer):
    employer_profile = EmployerProfileSerializer(read_only=True)
    applicant_profile = ApplicantProfileSerializer(read_only=True)
    
    class Meta:
        model = User
        fields = ["id", "email", "first_name", "last_name", "role", "firebase_uid", "employer_profile", "applicant_profile",]


class EmployerSignupSerializer(serializers.ModelSerializer):
    company_name = serializers.CharField(required=True)
    job_title = serializers.CharField(required=True)
    company_size = serializers.CharField(required=False, allow_blank=True)
    company_website = serializers.URLField(required=False, allow_blank=True)

    class Meta:
        model = User
        fields = [
            "email", "password", "first_name", "last_name", "role",
            "company_name", "job_title", "company_size", "company_website"
        ]
        extra_kwargs = {"password": {"write_only": True}}

    def create(self, validated_data):
        password = validated_data.pop("password")
        company_name = validated_data.pop("company_name")
        job_title = validated_data.pop("job_title")
        company_size = validated_data.pop("company_size", "")
        company_website = validated_data.pop("company_website", "")

        # create or get company
        company, created = Company.objects.get_or_create(
            name__iexact=company_name,  # case-insensitive match
            defaults={
                "name": company_name,
                "size": company_size,
                "website": company_website,
            }
        )

        # create user
        user = User.objects.create(**validated_data)
        user.set_password(password)
        user.save()

        EmployerProfile.objects.create(
            user=user,
            company=company,
            job_title=job_title,
            # company_name=company_name,
            # job_title=job_title,
            # company_size=company_size,
            # company_website=company_website,
        )
        return user

def extract_text_from_resume(file_field) -> str:
    # file_field is an InMemoryUploadedFile
    file_path = file_field.temporary_file_path() if hasattr(file_field, "temporary_file_path") else None
    temp_file_created = False
    
    try:
        if not file_path:
            # handle in-memory files
            file_path = f"temp_resume_{file_field.name}"
            with open(file_path, "wb") as f:
                for chunk in file_field.chunks():
                    f.write(chunk)
            temp_file_created = True

        mime_type, _ = mimetypes.guess_type(file_field.name)

        if mime_type == "application/pdf":
            text = ""
            with fitz.open(file_path) as pdf:
                for page in pdf:
                    text += page.get_text()
            return text

        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)

        elif mime_type == "application/msword":
            converted = file_path + ".docx"
            result = subprocess.run(["libreoffice", "--headless", "--convert-to", "docx", file_path], 
                                  capture_output=True, text=True)
            if result.returncode != 0:
                raise ValueError(f"Failed to convert .doc file: {result.stderr}")
            doc = Document(converted)
            os.remove(converted)
            return "\n".join(p.text for p in doc.paragraphs)

        else:
            raise ValueError(f"Unsupported file type: {mime_type}")
    
    finally:
        # Clean up temporary file if we created one
        if temp_file_created and os.path.exists(file_path):
            os.remove(file_path)

class ApplicantSignupSerializer(serializers.ModelSerializer):
    major = serializers.CharField(required=True)
    school = serializers.CharField(required=True)
    resume = serializers.CharField(required=False, allow_blank=True)
    resume_file = serializers.FileField(required=False, allow_null=True)
    skills = serializers.CharField(required=False, allow_blank=True)
    portfolio_url = serializers.URLField(required=False, allow_blank=True)

    class Meta:
        model = User
        fields = [
            "email", "password", "first_name", "last_name", "role",
            "major", "school", "resume", "resume_file", "skills", "portfolio_url"
        ]
        extra_kwargs = {"password": {"write_only": True}}

    def create(self, validated_data):
        password = validated_data.pop("password")
        major = validated_data.pop("major")
        school = validated_data.pop("school")
        resume = validated_data.pop("resume", "")
        resume_file = validated_data.pop("resume_file", None)
        skills = validated_data.pop("skills", "")
        portfolio_url = validated_data.pop("portfolio_url", "")
        vector_embedding = None

        if resume_file:
            try:
                text = extract_text_from_resume(resume_file)
                if text.strip():  # Only process if text was extracted
                    vector_new = embedding_model.encode([text])[0]
                    # Normalize the vector to unit length
                    norm = np.linalg.norm(vector_new)
                    if norm > 0:
                        vector_embedding = (vector_new / norm).tolist()
                    else:
                        vector_embedding = vector_new.tolist()
            except Exception as e:
                # Log the error but don't fail the user creation
                print(f"Error processing resume for vector embedding: {e}")
                vector_embedding = None

        user = User.objects.create(**validated_data)
        user.set_password(password)
        user.save()

        ApplicantProfile.objects.create(
            user=user,
            major=major,
            school=school,
            resume=resume,
            resume_file=resume_file,
            skills=skills,
            portfolio_url=portfolio_url,
            vector_embedding=vector_embedding
        )
        return user
    